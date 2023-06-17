from docx import Document

def caption_to_md(cap_filename : str, md_filename : str) -> None:
    '''
    convert a caption file to a md file
    
    Args:
        cap_filename : the name of the caption file ending with .srt
        md_filename : the name of the md file to be saved ending with .md
    
    Return:
        None
    '''
    #read the caption file
    #get id and caption
    '''
    1
    00:00:00,000 --> 00:00:03,832
    最近天气好好，想出去拍照片

    2
    00:00:03,832 --> 00:00:07,808
    我也有这个打算，不过我的相机是渣渣

    3
    00:00:07,808 --> 00:00:13,272
    哈哈哈我也不是专业的，我就是瞎拍，多拍拍就好了
    '''
    cap={}
    with open(cap_filename, "r", encoding='utf-8') as f:
        for line in f:
            if line.strip().isdigit():
                id = line.strip()
                cap[id] = ""
            elif "-->" in line:
                continue
            elif line.strip() == "":
                continue
            else:
                cap[id] += line.strip()
    #write to md file
    '''
    # id
    caption
    '''
    with open(md_filename, "w", encoding='utf-8') as f:
        f.write("# Caption\n\n")
        for id in cap:
            f.write(f"## {id}\n\n")
            f.write(f"{cap[id]}\n\n")

def caption_to_docx(cap_filename : str, docx_filename : str) -> None:
    '''
    convert a md file to a docx file
    
    Args:
        cap_filename : the name of the caption file ending with .srt
        docx_filename : the name of the docx file to be saved ending with .docx
        
    Return:
        None
    '''
    #read the caption file
    #get id and caption
    cap={}
    with open(cap_filename, "r", encoding='utf-8') as f:
        for line in f:
            if line.strip().isdigit():
                id = line.strip()
                cap[id] = ""
            elif "-->" in line:
                continue
            elif line.strip() == "":
                continue
            else:
                cap[id] += line.strip()
    #write to docx file
    
    """
    id -> style: bold,font size: 24
    caption -> style: normal,font size: 24


    """

    document = Document()
    for id in cap:
        document.add_heading(id, level=2)
        document.add_paragraph(cap[id])
    document.save(docx_filename)

    

if __name__ == "__main__":
    caption_to_md("output.srt", "test.md")
    caption_to_docx("output.srt", "test.docx")
